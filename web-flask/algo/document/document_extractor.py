
"""
文档内容提取工具
支持txt、docx、pdf、doc等格式的文档内容提取
"""
import os
import re
from typing import Dict, Optional
import docx
import pdfplumber

class DocumentExtractor:
    """文档内容提取器"""
    
    @staticmethod
    def extract_content(file_path: str, file_type: str) -> Dict:
        """
        提取文档内容
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            
        Returns:
            Dict: 包含提取结果的字典
        """
        try:
            if file_type.lower() == 'txt':
                return DocumentExtractor._extract_txt(file_path)
            elif file_type.lower() == 'docx':
                return DocumentExtractor._extract_docx(file_path)
            elif file_type.lower() == 'doc':
                return DocumentExtractor._extract_doc(file_path)
            elif file_type.lower() == 'pdf':
                return DocumentExtractor._extract_pdf(file_path)
            else:
                return {
                    'success': False,
                    'error': f'不支持的文件类型: {file_type}'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'文档内容提取失败: {str(e)}'
            }
    
    @staticmethod
    def _extract_txt(file_path: str) -> Dict:
        """提取TXT文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 清理内容
            content = DocumentExtractor._clean_content(content)
            
            return {
                'success': True,
                'content': content,
                'summary': DocumentExtractor._generate_summary(content)
            }
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
                content = DocumentExtractor._clean_content(content)
                return {
                    'success': True,
                    'content': content,
                    'summary': DocumentExtractor._generate_summary(content)
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'TXT文件读取失败: {str(e)}'
                }
    
    @staticmethod
    def _extract_doc(file_path: str) -> Dict:
        """
        提取DOC文件内容
        注意：.doc是旧版Word格式，建议转换为.docx格式
        """
        return {
            'success': False,
            'error': '不支持.doc格式的文件。请将文件另存为.docx格式后重新上传。'
        }
    
    @staticmethod
    def _extract_docx(file_path: str) -> Dict:
        """
        提取DOCX文件内容
        """
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(' | '.join(row_text))
            
            content = '\n'.join(content_parts)
            content = DocumentExtractor._clean_content(content)
            
            return {
                'success': True,
                'content': content,
                'summary': DocumentExtractor._generate_summary(content)
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX文件读取失败: {str(e)}'
            }
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Dict:
        """使用pdfplumber提取PDF文件内容"""
        try:
            content_parts = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        content_parts.append(text.strip())

            content = '\n'.join(content_parts).strip()
            content = DocumentExtractor._clean_content(content)

            return {
                'success': True,
                'content': content,
                'summary': DocumentExtractor._generate_summary(content)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'PDF文件读取失败: {str(e)}'
            }
    
    @staticmethod
    def _clean_content(content: str) -> str:
        """清理文档内容，避免误删正常符号"""
        if not content:
            return ""

        # 去掉控制字符
        content = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', content)

        # 统一空白符
        content = re.sub(r'[ \t]+', ' ', content)
        content = re.sub(r'\n\s*\n+', '\n', content)

        return content.strip()
    
    @staticmethod
    def _generate_summary(content: str, max_length: int = 200) -> str:
        """生成文档摘要"""
        if not content:
            return ""
        
        # 简单摘要：取前200个字符
        summary = content[:max_length]
        if len(content) > max_length:
            summary += "..."
        
        return summary
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """根据文件名获取文件类型"""
        ext = os.path.splitext(filename)[1].lower()
        type_map = {
            '.txt': 'txt',
            '.docx': 'docx',
            '.doc': 'doc',
            '.pdf': 'pdf'
        }
        return type_map.get(ext, 'unknown')
    
    @staticmethod
    def is_supported_file_type(filename: str) -> bool:
        """检查是否为支持的文件类型"""
        file_type = DocumentExtractor.get_file_type(filename)
        return file_type in ['txt', 'docx', 'pdf']
